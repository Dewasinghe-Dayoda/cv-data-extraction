import json
import re
import time
import os
from datetime import datetime, timezone
import pdfplumber
from config import (
    GEMINI_API_KEY, GEMINI_MODEL,
    GROQ_API_KEY, GROQ_MODEL,
    AI_PROVIDER, EXTRACTION_PROMPT
)

MONTH_MAP = {
    "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
    "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6,
    "jul": 7, "july": 7, "aug": 8, "august": 8, "sep": 9, "september": 9,
    "oct": 10, "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12
}

PRESENT_KEYWORDS = ["present", "current", "till date", "till now", "ongoing", "now"]


def _parse_date(date_str: str) -> tuple[int, int] | None:
    date_str = date_str.strip().lower().rstrip(".")
    for kw in PRESENT_KEYWORDS:
        if kw in date_str:
            now = datetime.now()
            return (now.year, now.month)
    m = re.match(r"(\d{4})", date_str)
    if m:
        return (int(m.group(1)), 1)
    m = re.match(r"([a-z]+)\s+(\d{4})", date_str)
    if m:
        mon = MONTH_MAP.get(m.group(1))
        if mon:
            return (int(m.group(2)), mon)
    m = re.match(r"(\d{1,2})[/\-](\d{4})", date_str)
    if m:
        return (int(m.group(2)), int(m.group(1)))
    m = re.match(r"([a-z]+)\s+(\d{1,2})[/\-,\s]+(\d{4})", date_str)
    if m:
        mon = MONTH_MAP.get(m.group(1))
        if mon:
            return (int(m.group(3)), mon)
    m = re.match(r"(\d{1,2})\s+([a-z]+)\s+(\d{4})", date_str)
    if m:
        mon = MONTH_MAP.get(m.group(2))
        if mon:
            return (int(m.group(3)), mon)
    return None


def _months_between(start: tuple[int, int], end: tuple[int, int]) -> int:
    return (end[0] - start[0]) * 12 + (end[1] - start[1])


def _format_months(total_months: int) -> str:
    years = total_months // 12
    months = total_months % 12
    if years > 0 and months > 0:
        y_label = "Year" if years == 1 else "Years"
        m_label = "Month" if months == 1 else "Months"
        return f"{years} {y_label} {months} {m_label}"
    elif years > 0:
        y_label = "Year" if years == 1 else "Years"
        return f"{years} {y_label}"
    else:
        m_label = "Month" if months == 1 else "Months"
        return f"{months} {m_label}"


EXPERIENCE_HEADER_PATTERN = re.compile(
    r"(?:^|\n)[ \t]*(?:PROFESSIONAL[ \t]+|WORK[ \t]+|EMPLOYMENT[ \t]+|CAREER[ \t]+)?EXPERIENCE\b"
    r"|(?:^|\n)[ \t]*EMPLOYMENT[ \t]+HISTORY\b"
    r"|(?:^|\n)[ \t]*CAREER[ \t]+HISTORY\b",
    re.IGNORECASE
)

NEXT_SECTION_PATTERN = re.compile(
    r"(?:^|\n)[ \t]*EDUCATION\b|"
    r"(?:^|\n)[ \t]*PROFESSIONAL[ \t]+QUALIFICATION"
    r"|(?:^|\n)[ \t]*SKILLS?\b|"
    r"(?:^|\n)[ \t]*CERTIFICATIONS?\b|"
    r"(?:^|\n)[ \t]*TRAININGS?\b|"
    r"(?:^|\n)[ \t]*ACHIEVEMENTS?\b|"
    r"(?:^|\n)[ \t]*AWARDS?\b|"
    r"(?:^|\n)[ \t]*LANGUAGES?\b|"
    r"(?:^|\n)[ \t]*REFERENCES?\b|"
    r"(?:^|\n)[ \t]*PERSONAL[ \t]+DETAILS?\b|"
    r"(?:^|\n)[ \t]*CONTACT\b|"
    r"(?:^|\n)[ \t]*SUMMARY\b|"
    r"(?:^|\n)[ \t]*PROFILE\b|"
    r"(?:^|\n)[ \t]*ABOUT\b",
    re.IGNORECASE
)


def _extract_experience_section(cv_text: str) -> str | None:
    header_match = EXPERIENCE_HEADER_PATTERN.search(cv_text)
    if not header_match:
        return None
    section_start = header_match.end()
    next_section = NEXT_SECTION_PATTERN.search(cv_text, section_start)
    section_end = next_section.start() if next_section else len(cv_text)
    return cv_text[section_start:section_end]


def calculate_experience_from_durations(cv_text: str) -> str | None:
    search_text = _extract_experience_section(cv_text)
    if search_text is None:
        return None

    duration_pattern = re.compile(
        r"(?<!\w)(\d+)\s+years?\b"
        r"(?:\s+(\d+)\s+months?\b)?"
        r"|(?<!\w)(\d+)\s+months?\b",
        re.IGNORECASE
    )

    total_months = 0
    for match in duration_pattern.finditer(search_text):
        if match.group(1):
            total_months += int(match.group(1)) * 12
            if match.group(2):
                total_months += int(match.group(2))
        elif match.group(3):
            total_months += int(match.group(3))

    if total_months <= 0:
        return None

    return _format_months(total_months)


def calculate_experience_from_text(cv_text: str) -> str | None:
    duration_result = calculate_experience_from_durations(cv_text)
    if duration_result:
        return duration_result

    date_range_pattern = re.compile(
        r"(?:^|(?<=\s)|(?<=\())(\d{4}|[a-zA-Z]{3,9} \d{4})[ \t]*"
        r"[\-–—to]+[ \t]*"
        r"(\d{4}|[a-zA-Z]{3,9} \d{4}|present|current|till[ \t]+(?:date|now))",
        re.IGNORECASE
    )

    education_degree_pattern = re.compile(
        r"\b(?:b\.?s\.?c\.?\b|b\.?a\.?\b|b\.?eng\b|b\.?tech\b|b\.?com\b|b\.?ed\b|b\.?sc\b|"
        r"m\.?a\.?\b|m\.?s\.?c\.?\b|m\.?tech\b|\bmba\b|m\.?com\b|m\.?ed\b|m\.?eng\b|"
        r"ph\.?d\b|d\.?phil\b|\bbachelor\b|\bmaster\b|\bdegree\b|\bdiploma\b|"
        r"\buniversity\b|\bcollege\b|\binstitut)",
        re.IGNORECASE
    )

    if not EXPERIENCE_HEADER_PATTERN.search(cv_text):
        return None

    periods = []
    for match in date_range_pattern.finditer(cv_text):
        start = _parse_date(match.group(1))
        end = _parse_date(match.group(2))
        if not start or not end:
            continue
        if _months_between(start, end) <= 0:
            continue

        match_start = match.start()
        match_end = match.end()

        line_start = cv_text.rfind("\n", 0, match_start)
        line_start = 0 if line_start == -1 else line_start + 1
        line_end = cv_text.find("\n", match_end)
        if line_end == -1:
            line_end = len(cv_text)
        same_line = cv_text[line_start:line_end]

        next_line_end = cv_text.find("\n", line_end + 1)
        if next_line_end == -1:
            next_line_end = len(cv_text)
        next_line = cv_text[line_end + 1:next_line_end] if line_end + 1 < len(cv_text) else ""

        if education_degree_pattern.search(same_line) or education_degree_pattern.search(next_line):
            continue

        periods.append((start, end))

    if not periods:
        return calculate_experience_from_durations(cv_text)

    periods.sort(key=lambda x: x[0])

    merged = [periods[0]]
    for start, end in periods[1:]:
        prev_start, prev_end = merged[-1]
        if start <= prev_end:
            if end > prev_end:
                merged[-1] = (prev_start, end)
        else:
            merged.append((start, end))

    total_months = sum(_months_between(s, e) for s, e in merged)

    if total_months <= 0:
        return None

    return _format_months(total_months)


class RateLimitError(Exception):
    def __init__(self, provider: str, message: str, retry_after: int = 0, is_daily_limit: bool = None):
        self.provider = provider
        self.retry_after = retry_after
        if is_daily_limit is not None:
            self.is_daily_limit = is_daily_limit
        else:
            self.is_daily_limit = "per day" in message.lower() or "tpd" in message.lower() or "rpd" in message.lower()
        super().__init__(message)


def extract_text_from_pdf(pdf_path: str) -> str:
    text_parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(docx_path: str) -> str:
    from docx import Document
    doc = Document(docx_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n\n".join(text_parts)


def extract_text_from_image(image_path: str) -> str:
    import pytesseract
    from PIL import Image
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    return text


def _parse_retry_after(error_msg: str) -> int:
    match = re.search(r"try again in (\d+)m(\d+\.?\d*)s", error_msg)
    if match:
        minutes = int(match.group(1))
        seconds = float(match.group(2))
        return int(minutes * 60 + seconds)
    match = re.search(r"retry-after[:\s]+(\d+)", error_msg, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return 0


def _extract_with_gemini(cv_text: str) -> str:
    from google import genai
    client = genai.Client(api_key=GEMINI_API_KEY)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    prompt = EXTRACTION_PROMPT.format(text=cv_text, today=today)
    max_retries = 2
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt
            )
            return response.text.strip()
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                retry_after = _parse_retry_after(error_msg)
                raise RateLimitError("Gemini", error_msg, retry_after)
            raise Exception(f"Gemini API error: {error_msg}")
    return ""


def _extract_with_groq(cv_text: str) -> str:
    from groq import Groq
    client = Groq(api_key=GROQ_API_KEY)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    prompt = EXTRACTION_PROMPT.format(text=cv_text, today=today)
    max_retries = 2
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": "You are a CV/resume data extraction assistant. You MUST return ONLY a valid JSON object. No explanations, no markdown, no text before or after. Just the raw JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,
                max_tokens=2048
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "rate_limit" in error_msg.lower():
                retry_after = _parse_retry_after(error_msg)
                raise RateLimitError("Groq", error_msg, retry_after)
            raise Exception(f"Groq API error: {error_msg}")
    return ""


def _parse_json_response(raw_response: str) -> dict:
    if raw_response.startswith("```"):
        raw_response = raw_response.split("\n", 1)[1]
        if raw_response.endswith("```"):
            raw_response = raw_response[:-3]
        raw_response = raw_response.strip()

    try:
        data = json.loads(raw_response)
    except json.JSONDecodeError:
        data = {}

    required_keys = [
        "candidate_name", "contact_number", "email", "gender", "age",
        "total_experience_years", "last_employer", "key_skills", "current_job_title"
    ]
    for key in required_keys:
        if key not in data:
            data[key] = None

    return data


def extract_candidate_data(cv_text: str) -> dict:
    errors = []
    rate_limit_errors = []

    providers = []
    if AI_PROVIDER == "groq":
        providers = ["groq"]
    elif AI_PROVIDER == "gemini":
        providers = ["gemini"]
    else:
        if GROQ_API_KEY:
            providers.append("groq")
        if GEMINI_API_KEY:
            providers.append("gemini")

    for provider in providers:
        try:
            if provider == "groq":
                raw_response = _extract_with_groq(cv_text)
            else:
                raw_response = _extract_with_gemini(cv_text)
            return _parse_json_response(raw_response)
        except RateLimitError as e:
            rate_limit_errors.append(e)
            errors.append(f"{e.provider}: {e}")
            continue
        except Exception as e:
            errors.append(f"{provider}: {e}")
            continue

    if rate_limit_errors:
        worst = max(rate_limit_errors, key=lambda x: x.retry_after)
        raise RateLimitError(
            worst.provider,
            f"Daily limit reached for {worst.provider}. Resets at midnight UTC.",
            worst.retry_after,
            is_daily_limit=True
        )

    raise Exception(f"All AI providers failed:\n" + "\n".join(errors))


def process_cv(pdf_path: str) -> tuple[dict, str]:
    ext = os.path.splitext(pdf_path)[1].lower()

    if ext == ".pdf":
        cv_text = extract_text_from_pdf(pdf_path)
    elif ext == ".docx":
        cv_text = extract_text_from_docx(pdf_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        cv_text = extract_text_from_image(pdf_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not cv_text.strip():
        raise ValueError(f"Could not extract text from {ext} file. The file may be empty or unreadable.")

    extracted_data = extract_candidate_data(cv_text)

    calculated_exp = calculate_experience_from_text(cv_text)
    if calculated_exp:
        extracted_data["total_experience_years"] = calculated_exp

    return extracted_data, cv_text
